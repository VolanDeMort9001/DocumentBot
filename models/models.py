from docx import Document
from docx.shared import Pt, Inches
from typing import Iterable


def create_document(text: Iterable[str]) -> None:
    #Создание документа
    doc = Document()
    print(text)
    doc.add_paragraph("Коллеги, добрый день!\n\nБлагодарим за вовремя присланный мониторинговый отчет.\nПо результатам его проверки просим учесть следующие рекомендации и доработать отчет в срок до:\n")
    number = 1
    for par in text:
        doc.add_paragraph(f"{number}. {par}")
        doc.add_paragraph("")
        number += 1
    doc.add_paragraph("Ждем доработанный отчет. Если возникнут вопросы, я на связи. Мои контакты отображены на площадке. Буду рада помочь.")
    doc.save("ABS_отчет.docx")